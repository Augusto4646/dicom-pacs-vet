import json
import base64
import requests
from django.http import JsonResponse, HttpResponse, Http404
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth.decorators import login_required
from django.contrib.auth import login
from django.shortcuts import render, redirect, get_object_or_404
from django.core.files.base import ContentFile
from django.template.loader import render_to_string
from rest_framework import viewsets
from docx import Document
from django.http import FileResponse, Http404
from django.contrib.auth.decorators import login_required
from .models import Modelo
from io import BytesIO
from PIL import Image
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import requests

from .models import (
    Usuario, Exame, Modelo, Laudo, Paciente, Instituicao,
    Financeiro, Clinica, VeterinarioPedidor, LancamentoFinanceiro
)
from .serializers import ExameSerializer
from .util import gerar_codigo_unico
from base.services.orthanc_sync import sincronizar_estudos
from base.models import Exame
from django.db.models import Sum
import io
import zipfile

ORTHANC_URL = "https://pacsvisionxvet.conexao46.com.br/orthanc"


# ---------------------------------------------------------------------------
# Sync
# ---------------------------------------------------------------------------

@login_required
def sync_exames(request):
    sincronizar_estudos()
    return redirect("dashbord")


# ---------------------------------------------------------------------------
# Webhook — recebe novo exame do Orthanc
# ---------------------------------------------------------------------------

@csrf_exempt
def novo_exame(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            print("Webhook recebido:", data)
            return JsonResponse({"status": "ok"})
        except Exception as e:
            print("Erro:", str(e))
            return JsonResponse({"error": str(e)}, status=500)
    return JsonResponse({"error": "invalid method"}, status=405)


# ---------------------------------------------------------------------------
# Financeiro
# ---------------------------------------------------------------------------

@login_required
def financeiro(request):
    if request.method == "POST":
        LancamentoFinanceiro.objects.create(
            data=request.POST.get("data"),
            paciente=request.POST.get("paciente"),
            tutor=request.POST.get("tutor"),
            clinica=request.POST.get("clinica"),
            tipo_exame=request.POST.get("tipo_exame"),
            valor=request.POST.get("valor"),
            forma_pagamento=request.POST.get("forma_pagamento"),
            pago=request.POST.get("pago") == "on",
            enviado=request.POST.get("enviado") == "on",
            observacoes=request.POST.get("observacoes"),
        )
        return redirect("financeiro")

    lancamentos    = LancamentoFinanceiro.objects.all().order_by("-data")
    total          = lancamentos.aggregate(total=Sum("valor"))["total"] or 0
    total_pago     = lancamentos.filter(pago=True).aggregate(total=Sum("valor"))["total"] or 0
    total_pendente = lancamentos.filter(pago=False).aggregate(total=Sum("valor"))["total"] or 0

    return render(request, "financeiro.html", {
        "lancamentos": lancamentos,
        "total": total,
        "total_pago": total_pago,
        "total_pendente": total_pendente,
    })


# ---------------------------------------------------------------------------
# Dashboard
# ---------------------------------------------------------------------------

@login_required
def dashbord(request):
    usuario, created = Usuario.objects.get_or_create(user=request.user)
    exames = Exame.objects.filter(
        instituicao__in=usuario.instituicoes.all()
    ).order_by('-id')
    return render(request, "dashbord.html", {"exames": exames})


@login_required
def menu(request):
    usuario, _ = Usuario.objects.get_or_create(user=request.user)
    inst   = usuario.instituicao_pertencente
    exames = Exame.objects.filter(instituicao__in=usuario.instituicoes.all()).order_by('-id')

    financeiros   = Financeiro.objects.filter(instituicao=inst)
    a_receber     = financeiros.filter(pago=False).aggregate(t=Sum('valor'))['t'] or 0
    recebido      = financeiros.filter(pago=True).aggregate(t=Sum('valor'))['t'] or 0
    despesa       = financeiros.aggregate(t=Sum('despesa'))['t'] or 0
    saldo_liquido = recebido - despesa

    total_aguardando = exames.filter(status='Aguardando').count()
    total_urgente    = exames.filter(status='Urgente').count()
    total_laudado    = exames.filter(status='Laudado').count()

    clinicas_devedoras = (
        Financeiro.objects
        .filter(pago=False, clinica__isnull=False)
        .values('clinica__nome_clinica', 'clinica__id')
        .annotate(total=Sum('valor'))
        .order_by('-total')
    )

    return render(request, "menu.html", {
        "a_receber": a_receber,
        "recebido": recebido,
        "despesa": despesa,
        "saldo_liquido": saldo_liquido,
        "exames_recentes": exames[:5],
        "total_aguardando": total_aguardando,
        "total_urgente": total_urgente,
        "total_laudado": total_laudado,
        "total_exames": exames.count(),
        "clinicas_devedoras": clinicas_devedoras,
    })


# ---------------------------------------------------------------------------
# Home / Login
# ---------------------------------------------------------------------------

def home(request):
    form = AuthenticationForm()
    return render(request, "home.html", {"form": form})


def login_view(request):
    if request.method == "POST":
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect("menu")
    else:
        form = AuthenticationForm()
    return render(request, "login.html", {"form": form})


def pagina_laudos(request):
    return render(request, "pagina_laudos.html", {})


# ---------------------------------------------------------------------------
# Salvar laudo (HTML)
# ---------------------------------------------------------------------------

@login_required
@csrf_exempt
def salvar_laudo(request):
    if request.method == "POST":
        data  = json.loads(request.body)
        exame = get_object_or_404(Exame, study_instance_uid=data["studyInstanceUid"])
        exame.laudo_html = data["laudoHTML"]
        exame.save()
        return JsonResponse({"codigo": exame.codigo_acesso})
    return JsonResponse({"error": "invalid method"}, status=405)


# ---------------------------------------------------------------------------
# Salvar PDF
# ---------------------------------------------------------------------------

@login_required
@csrf_exempt
def salvar_pdf(request, exame_id):
    if request.method == "POST":
        exame = get_object_or_404(Exame, id=exame_id)
        data  = request.POST.get("pdf_base64")
        if not data:
            return JsonResponse({"error": "pdf_base64 ausente"}, status=400)
        format, imgstr = data.split(";base64,")
        pdf_bytes = base64.b64decode(imgstr)
        exame.pdf.save(f"laudo_{exame.id}.pdf", ContentFile(pdf_bytes), save=True)
        return JsonResponse({"status": "ok"})
    return JsonResponse({"error": "invalid method"}, status=405)


# ---------------------------------------------------------------------------
# Upload de PDF via formulário
# ---------------------------------------------------------------------------

@login_required
def upload_laudo_pdf(request):
    if request.method == "POST":
        exame_id = request.POST.get("exame_id")
        pdf      = request.FILES.get("pdf")
        exame    = get_object_or_404(Exame, id=exame_id)
        exame.laudo_pdf.save(f"laudo_{exame.id}.pdf", pdf)
        return JsonResponse({"ok": True})
    return JsonResponse({"error": "invalid method"}, status=405)


# ---------------------------------------------------------------------------
# Portal do paciente
# ---------------------------------------------------------------------------

def portal_exame(request, codigo=None):
    if codigo:
        exame = Exame.objects.filter(codigo_acesso=codigo).first()
        if exame:
            return render(request, "portal_exame.html", {"exame": exame})
        return render(request, "portal.html", {"erro": "Código não encontrado"})
    return render(request, "portal.html")


# ---------------------------------------------------------------------------
# Editar nome do paciente
# ---------------------------------------------------------------------------

@login_required
def editar_paciente(request, exame_id):
    if request.method == "POST":
        novo_nome = request.POST.get("nome")
        exame     = get_object_or_404(Exame, id=exame_id)
        study_uid = exame.study_instance_uid

        find = requests.post(
            f"{ORTHANC_URL}/tools/find",
            json={"Level": "Study", "Query": {"StudyInstanceUID": study_uid}}
        ).json()

        if not find:
            return JsonResponse({"erro": "Estudo não encontrado no Orthanc"})

        requests.post(
            f"{ORTHANC_URL}/studies/{find[0]}/modify",
            json={"Replace": {"PatientName": novo_nome}, "KeepSource": False}
        )

        exame.paciente.nome = novo_nome
        exame.paciente.save()
        return JsonResponse({"ok": True})

    return JsonResponse({"error": "invalid method"}, status=405)


# ---------------------------------------------------------------------------
# Baixar DICOM como ZIP
# ---------------------------------------------------------------------------

def baixar_dicom(request, exame_id):
    try:
        exame = Exame.objects.get(id=exame_id)
    except Exame.DoesNotExist:
        raise Http404("Exame não encontrado")

    if not exame.orthanc_ids:
        return HttpResponse("Sem DICOM disponível", status=400)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        for i, orthanc_id in enumerate(exame.orthanc_ids):
            url      = f"{ORTHANC_URL}/instances/{orthanc_id}/file"
            response = requests.get(url)
            if response.status_code == 200:
                zip_file.writestr(f"exame_{exame.id}_{i+1}.dcm", response.content)

    zip_buffer.seek(0)
    http_response = HttpResponse(zip_buffer, content_type="application/zip")
    http_response["Content-Disposition"] = f'attachment; filename="exame_{exame.id}.zip"'
    return http_response


# ---------------------------------------------------------------------------
# DICOM → JPG para o viewer USG
# ---------------------------------------------------------------------------

@login_required
def dicom_para_imagens(request, exame_id):
    exame = get_object_or_404(Exame, id=exame_id)

    try:
        r = requests.post(
            f"{ORTHANC_URL}/tools/find",
            json={
                "Level": "Study",
                "Query": {"StudyInstanceUID": exame.study_instance_uid},
                "Expand": False,
            },
            timeout=10,
        )

        print("STATUS:", r.status_code)
        print("BODY:", r.text[:1000])

        results  = r.json()
        study_id = results[0] if results else None

        print("study_id =", study_id)
        print("uid =", exame.study_instance_uid)

    except Exception as e:
        print(f"[dicom_para_imagens] Erro ao buscar study: {e}")
        return JsonResponse({"imagens": [], "erro": str(e)}, status=400)

    if not study_id:
        return JsonResponse(
            {"imagens": [], "erro": "Study não encontrado no Orthanc"},
            status=400
        )

    from .dicom_utils import get_jpg_urls_for_study

    print("USANDO UID:", exame.study_instance_uid)
    print("USANDO ORTHANC ID:", study_id)

    urls = get_jpg_urls_for_study(exame.study_instance_uid)

    print("URLS:", urls)

    return JsonResponse({"imagens": urls})


# ---------------------------------------------------------------------------
# Modelos de laudo
# ---------------------------------------------------------------------------

def listar_modelos(request):
    usuario, created = Usuario.objects.get_or_create(user=request.user)
    modelos = Modelo.objects.all()
    return render(request, "pagina_modelos.html", {"modelos": modelos})


@login_required
def criar_modelos(request):
    if request.method == "POST":
        usuario, _ = Usuario.objects.get_or_create(user=request.user)

        modelo = Modelo.objects.create(
            usuario_logado=usuario,
            nome_modelo=request.POST.get("nome_modelo"),
            campo1=request.POST.get("campo1"),
            campo2=request.POST.get("campo2"),
            campo3=request.POST.get("campo3"),
            campo4=request.POST.get("campo4"),
            campo5=request.POST.get("campo5"),
            campo6=request.POST.get("campo6"),
            campo7=request.POST.get("campo7"),
            campo8=request.POST.get("campo8"),
            campo9=request.POST.get("campo9"),
            campo10=request.POST.get("campo10"),
            campo11=request.POST.get("campo11"),
            campo12=request.POST.get("campo12"),
        )

        if request.FILES.get("arquivo_docx"):
            modelo.arquivo_docx = request.FILES["arquivo_docx"]
            modelo.save()

        return redirect("/listar_modelos/")

    return redirect("/listar_modelos/")


def deletar_modelo(request, modelo_id):
    modelo = get_object_or_404(Modelo, id=modelo_id)
    modelo.delete()
    return redirect('/listar_modelos/')


# ---------------------------------------------------------------------------
# Clínicas
# ---------------------------------------------------------------------------

def listar_clinicas(request):
    usuario, _ = Usuario.objects.get_or_create(user=request.user)
    clinicas   = Clinica.objects.all()
    return render(request, "pagina_clinicas.html", {"clinicas": clinicas})


def criar_clinica(request):
    if request.method == "POST":
        usuario, created = Usuario.objects.get_or_create(user=request.user)
        Clinica.objects.create(
            usuario_logado=usuario,
            nome_clinica=request.POST.get("nome_clinica"),
            whats_clinica=request.POST.get("whats_clinica"),
        )
        clinicas = Clinica.objects.filter(usuario_logado=usuario)
        return render(request, "pagina_clinicas.html", {"clinicas": clinicas})


def deletar_clinica(request, clinica_id):
    clinica = get_object_or_404(Clinica, id=clinica_id)
    clinica.delete()
    return redirect('/listar_clinicas/')


def editar_clinica(request, clinica_id):
    clinica = get_object_or_404(Clinica, id=clinica_id)
    if request.method == "POST":
        clinica.nome_clinica  = request.POST.get("nome_clinica")
        clinica.whats_clinica = request.POST.get("whats_clinica")
        clinica.save()
        return redirect('/listar_clinicas/')
    return render(request, "editar_clinica.html", {"clinica": clinica})


# ---------------------------------------------------------------------------
# Lançamentos financeiros
# ---------------------------------------------------------------------------

def excluir_lancamento(request, id):
    LancamentoFinanceiro.objects.filter(pk=id).delete()
    return HttpResponse(status=200)


def importar_lancamentos(request):
    data = json.loads(request.body)
    objs = [LancamentoFinanceiro(**row) for row in data['lancamentos']]
    LancamentoFinanceiro.objects.bulk_create(objs, ignore_conflicts=True)
    return JsonResponse({'criados': len(objs), 'erros': 0})


# ---------------------------------------------------------------------------
# Editor de laudo
# ---------------------------------------------------------------------------

@login_required
def laudo_editor(request, exame_id):
    usuario, created = Usuario.objects.get_or_create(user=request.user)
    exame    = get_object_or_404(Exame, id=exame_id)
    exames   = Exame.objects.filter(instituicao__in=usuario.instituicoes.all())
    membros  = exame.instituicao.membros_insituticao.all() if exame.instituicao else []
    modelos  = Modelo.objects.all()
    clinicas = Clinica.objects.filter(usuario_logado=usuario)
    vets     = VeterinarioPedidor.objects.all()

    if request.method == 'POST':
        vet_id = request.POST.get("veterinario_pedidor")
        if vet_id:
            exame.veterinario_pedidor_id = vet_id

        exame.tipo_exame = request.POST.get("tipo_exame")
        if not exame.codigo_acesso:
            exame.codigo_acesso = exame._gerar_codigo()
        exame.save()

        exame.paciente.nome       = request.POST.get("Paciente")
        exame.paciente.nome_tutor = request.POST.get("Tutor")
        exame.paciente.save()

        valor      = request.POST.get("valor") or 0
        repasse    = request.POST.get("repasse") or 0
        pago       = request.POST.get("pago") == "on"
        forma      = request.POST.get("forma_pagamento") or ""
        clinica_id = request.POST.get("clinica") or None

        if exame.financeiro:
            exame.financeiro.pago                    = pago
            exame.financeiro.valor                   = valor
            exame.financeiro.valor_repasse_a_clinica = repasse
            exame.financeiro.forma_pagamento         = forma
            if clinica_id:
                exame.financeiro.clinica_id = clinica_id
            exame.financeiro.save()
        elif exame.instituicao:
            financeiro = Financeiro.objects.create(
                instituicao=exame.instituicao,
                pago=pago,
                valor=valor,
                valor_repasse_a_clinica=repasse,
                forma_pagamento=forma,
                clinica_id=clinica_id,
            )
            exame.financeiro = financeiro
            exame.save()

    return render(request, "laudo_editor.html", {
        "exame": exame,
        "exames": exames,
        "membros": membros,
        "modelos": modelos,
        "clinicas": clinicas,
        "veterinarios": vets,
    })


# ---------------------------------------------------------------------------
# Monitor (segundo monitor)
# ---------------------------------------------------------------------------

def view_monitor(request, exame_id):
    usuario, created = Usuario.objects.get_or_create(user=request.user)
    exame    = get_object_or_404(Exame, id=exame_id)
    exames   = Exame.objects.filter(instituicao__in=usuario.instituicoes.all())
    membros  = exame.instituicao.membros_insituticao.all() if exame.instituicao else []
    modelos  = Modelo.objects.all()
    clinicas = Clinica.objects.filter(usuario_logado=usuario)
    vets     = VeterinarioPedidor.objects.all()

    return render(request, "lado_esquerdo.html", {
        "exame": exame,
        "exames": exames,
        "membros": membros,
        "modelos": modelos,
        "clinicas": clinicas,
        "veterinarios": vets,
    })


# ---------------------------------------------------------------------------
# Atualizar status
# ---------------------------------------------------------------------------

def atualizar_status(request, exame_id):
    if request.method == "POST":
        dados = json.loads(request.body)
        Exame.objects.filter(id=exame_id).update(status=dados['status'])
        return JsonResponse({'ok': True})


def atualizar_status_laudo_editor(request, exame_id):
    if request.method == 'POST':
        Exame.objects.filter(id=exame_id).update(status="Laudado")
        return JsonResponse({'ok': True})


def dashboard_visual(request):
    exames  = Exame.objects.all()
    colunas = [
        {"status": "Aguardando", "tema": "sunset",  "titulo": "⏳ Aguardando"},
        {"status": "Urgente",    "tema": "abyss",   "titulo": "🚨 Urgente"},
        {"status": "Laudado",    "tema": "forest",  "titulo": "✅ Laudado"},
    ]
    return render(request, "dashbord.html", {"exames": exames, "colunas": colunas})


# ---------------------------------------------------------------------------
# Editar cabeçalho DICOM
# ---------------------------------------------------------------------------

def editar_cabecalho(request, exame_id):
    if request.method == "POST":
        novo_nome = request.POST.get("novo_nome")
        exame     = get_object_or_404(Exame, id=exame_id)
        for orthanc_id in exame.orthanc_ids:
            requests.post(
                f"{ORTHANC_URL}/instances/{orthanc_id}/modify",
                json={"Replace": {"PatientName": novo_nome}, "Force": True}
            )
    return JsonResponse({"status": "ok"})


# ---------------------------------------------------------------------------
# Veterinários
# ---------------------------------------------------------------------------

def listar_veterinarios(request):
    vets = VeterinarioPedidor.objects.all()
    return render(request, "pagina_veterinarios.html", {"veterinarios": vets})


def criar_veterinario(request):
    if request.method == "POST":
        VeterinarioPedidor.objects.create(nome=request.POST.get("nome"))
        vets = VeterinarioPedidor.objects.all()
        return render(request, "pagina_veterinarios.html", {"veterinarios": vets})


def deletar_veterinario(request, veterinario_id):
    veterinario = get_object_or_404(VeterinarioPedidor, id=veterinario_id)
    veterinario.delete()
    return redirect('/listar_veterinarios/')


# ---------------------------------------------------------------------------
# Deletar exame
# ---------------------------------------------------------------------------

def deletar_exame(request, exame_id):
    exame = get_object_or_404(Exame, id=exame_id)
    exame.delete()
    return redirect('/dashbord/')


# ---------------------------------------------------------------------------
# Atualizar tipo de exame
# ---------------------------------------------------------------------------

@login_required
def atualizar_tipo_exame(request, exame_id):
    if request.method == "POST":
        dados = json.loads(request.body)
        Exame.objects.filter(id=exame_id).update(tipo_exame=dados.get('tipo_exame'))
        return JsonResponse({'ok': True})
    return JsonResponse({'error': 'invalid method'}, status=405)


# ---------------------------------------------------------------------------
# Get modelo HTML
# ---------------------------------------------------------------------------

@login_required
def get_modelo_html(request, modelo_id, exame_id):
    from datetime import date

    modelo = get_object_or_404(Modelo, id=modelo_id)
    exame  = get_object_or_404(Exame, id=exame_id)

    html = modelo.html_conteudo or ''

    html = html.replace('{{paciente}}',    exame.paciente.nome or '')
    html = html.replace('{{tutor}}',       exame.paciente.nome_tutor or '')
    html = html.replace('{{raca}}',        exame.paciente.raca or '')
    html = html.replace('{{sexo}}',        exame.paciente.sexo or '')
    html = html.replace('{{idade}}',       str(exame.paciente.idade) if exame.paciente.idade else '')
    html = html.replace('{{veterinario}}', exame.veterinario_pedidor.nome if exame.veterinario_pedidor else '')
    html = html.replace(
        '{{clinica}}',
        exame.financeiro.clinica.nome_clinica
        if exame.financeiro and exame.financeiro.clinica
        else ''
    )
    html = html.replace('{{data}}', date.today().strftime('%d/%m/%Y'))

    return JsonResponse({'html': html})


# ---------------------------------------------------------------------------
# Blank DOCX para OnlyOffice
# ---------------------------------------------------------------------------

def blank_docx(request):
    from datetime import date

    modelo_id = request.GET.get("modelo")
    exame_id  = request.GET.get("exame_id")

    if not modelo_id:
        doc = Document()
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return HttpResponse(
            buffer.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    modelo = get_object_or_404(Modelo, id=modelo_id)

    if not modelo.arquivo_docx:
        doc = Document()
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return HttpResponse(
            buffer.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Monta substituições com dados do exame se disponível
    subs = {
        '{{paciente}}':    '',
        '{{tutor}}':       '',
        '{{raca}}':        '',
        '{{sexo}}':        '',
        '{{idade}}':       '',
        '{{veterinario}}': '',
        '{{clinica}}':     '',
        '{{data}}':        date.today().strftime('%d/%m/%Y'),
    }

    if exame_id:
        try:
            exame = Exame.objects.get(id=exame_id)
            subs['{{paciente}}']    = exame.paciente.nome or ''
            subs['{{tutor}}']       = exame.paciente.nome_tutor or ''
            subs['{{raca}}']        = exame.paciente.raca or ''
            subs['{{sexo}}']        = exame.paciente.sexo or ''
            subs['{{idade}}']       = str(exame.paciente.idade) if exame.paciente.idade else ''
            subs['{{veterinario}}'] = exame.veterinario_pedidor.nome if exame.veterinario_pedidor else ''
            subs['{{clinica}}']     = (
                exame.financeiro.clinica.nome_clinica
                if exame.financeiro and exame.financeiro.clinica else ''
            )
        except Exame.DoesNotExist:
            pass

    doc = Document(modelo.arquivo_docx.path)

    def substituir_paragrafo(paragrafo, subs):
        # Merge todos os runs em um texto único, substitui, redistribui
        texto_completo = ''.join(run.text for run in paragrafo.runs)
        modificado = False
        for chave, valor in subs.items():
            if chave in texto_completo:
                texto_completo = texto_completo.replace(chave, valor)
                modificado = True
        if modificado and paragrafo.runs:
            paragrafo.runs[0].text = texto_completo
            for run in paragrafo.runs[1:]:
                run.text = ''

    for para in doc.paragraphs:
        substituir_paragrafo(para, subs)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for para in celula.paragraphs:
                    substituir_paragrafo(para, subs)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response['Access-Control-Allow-Origin'] = '*'
    return response
    

def servir_docx_modelo(request, modelo_id):
    try:
        modelo = Modelo.objects.get(id=modelo_id)
    except Modelo.DoesNotExist:
        raise Http404

    if not modelo.arquivo_docx:
        raise Http404

    response = FileResponse(
        modelo.arquivo_docx.open('rb'),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'inline; filename="modelo_{modelo_id}.docx"'
    response['Access-Control-Allow-Origin'] = '*'
    return response


# ---------------------------------------------------------------------------
# Gerar PDF completo (laudo OnlyOffice + imagens DICOM)
# ---------------------------------------------------------------------------

@login_required
@csrf_exempt
def gerar_pdf_completo(request, exame_id):
    from .dicom_utils import get_jpg_urls_for_study

    exame = get_object_or_404(Exame, id=exame_id)

    if request.method != 'POST':
        return HttpResponse("Method not allowed", status=405)

    data    = json.loads(request.body)
    pdf_url = data.get('pdf_url')
    imagens = data.get('imagens', [])

    # 1. Baixa o PDF gerado pelo OnlyOffice
    pdf_bytes = requests.get(pdf_url, timeout=30).content
    writer    = PdfWriter()
    reader    = PdfReader(BytesIO(pdf_bytes))
    for page in reader.pages:
        writer.add_page(page)

    # 2. Anexa imagens DICOM
    BASE_URL = "https://pacsvisionxvet.conexao46.com.br"
    posicoes = [(20, 440), (300, 440), (20, 40), (300, 40)]

    for i in range(0, len(imagens), 4):
        packet = BytesIO()
        c      = canvas.Canvas(packet)
        grupo  = imagens[i:i+4]
        teve_imagem = False

        for idx, url in enumerate(grupo):
            if url.startswith('/'):
                url = BASE_URL + url
            try:
                img_bytes = requests.get(url, timeout=10).content
                img = Image.open(BytesIO(img_bytes))
                c.drawImage(
                    ImageReader(img),
                    posicoes[idx][0], posicoes[idx][1],
                    width=250, height=350,
                    preserveAspectRatio=True,
                )
                teve_imagem = True
            except Exception as e:
                print(f"[PDF] Erro imagem: {e}")

        if teve_imagem:
            c.showPage()
            c.save()
            packet.seek(0)
            pages = PdfReader(packet).pages
            if pages:
                writer.add_page(pages[0])

    output = BytesIO()
    writer.write(output)
    output.seek(0)

    # 3. Salva o PDF no exame para o portal
    exame.pdf.save(
        f"laudo_{exame.id}.pdf",
        ContentFile(output.getvalue()),
        save=True
    )
    output.seek(0)

    return HttpResponse(
        output.getvalue(),
        content_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="laudo_{exame.id}.pdf"'},
    )


# ---------------------------------------------------------------------------
# OnlyOffice callback — salva o DOCX gerado
# ---------------------------------------------------------------------------

@csrf_exempt
def onlyoffice_callback(request, exame_id):
    data = json.loads(request.body)
    print(data)

    if data.get("status") in [2, 4]:
        print("STATUS =", data.get("status"))
        print("URL =", data.get("url"))

        arquivo_url   = data.get("url")
        arquivo_bytes = requests.get(arquivo_url).content
        exame         = Exame.objects.get(id=exame_id)
        exame.docx.save(
            f"laudo_{exame.id}.docx",
            ContentFile(arquivo_bytes),
            save=True
        )
        print("DOCX SALVO")

    return JsonResponse({"error": 0})


# ---------------------------------------------------------------------------
# Forçar save no OnlyOffice
# ---------------------------------------------------------------------------

@csrf_exempt
def forcar_save_onlyoffice(request, exame_id):
    exame = get_object_or_404(Exame, id=exame_id)
    resp  = requests.post(
        "https://pacsvisionxvet.conexao46.com.br/onlyoffice/coauthoring/CommandService.ashx",
        json={"c": "forcesave", "key": f"laudo_{exame.id}"},
        timeout=10,
    )
    print("[forcesave] status:", resp.status_code, "body:", resp.text)
    return JsonResponse({"ok": True})


# ---------------------------------------------------------------------------
# Check DOCX
# ---------------------------------------------------------------------------

@login_required
def check_docx(request, exame_id):
    exame = get_object_or_404(Exame, id=exame_id)
    return JsonResponse({"ok": bool(exame.docx)})